from azure.identity import AzureCliCredential, InteractiveBrowserCredential
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import SearchIndex, SimpleField, SearchableField, SearchFieldDataType
from azure.storage.blob import BlobServiceClient
import sys


# PDF text extraction function
def extract_pdf_text(pdf_bytes):
    """Extract text content from PDF bytes using PyPDF2"""
    try:
        import PyPDF2
        import io

        pdf_file = io.BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        # Check if PDF is encrypted/protected
        if pdf_reader.is_encrypted:
            return "PDF_PROTECTED: This PDF document is password-protected or encrypted and cannot be processed."

        text_content = []
        for page in pdf_reader.pages:
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_content.append(page_text)
            except Exception:
                continue

        full_text = "\n".join(text_content).strip()

        # Check for protection messages
        protection_indicators = [
            "protected by Microsoft Office",
            "You'll need a different reader",
            "Download a compatible PDF reader",
            "This PDF Document has been protected"
        ]

        if any(indicator.lower() in full_text.lower() for indicator in protection_indicators):
            return "PDF_PROTECTED: This PDF document appears to be protected or encrypted."

        return full_text if full_text else "PDF_NO_TEXT: No readable text content found in PDF."

    except ImportError:
        return "PDF_ERROR: PyPDF2 library not available. Install with: pip install PyPDF2"
    except Exception as e:
        return f"PDF_ERROR: Error reading PDF content: {str(e)}"
    

# DOCX text extraction function
def extract_docx_text(docx_bytes):
    """Extract text content from DOCX bytes using python-docx"""
    try:
        from docx import Document
        import io

        docx_file = io.BytesIO(docx_bytes)
        doc = Document(docx_file)

        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)

        full_text = "\n".join(text_content).strip()
        return full_text if full_text else "DOCX_NO_TEXT: No readable text content found in DOCX."

    except ImportError:
        return "DOCX_ERROR: python-docx library not available. Install with: pip install python-docx"
    except Exception as e:
        return f"DOCX_ERROR: Error reading DOCX content: {str(e)}"

if len(sys.argv) < 4:
    print("Usage: python index_datasets.py <storage_account_name> <blob_container_name> <ai_search_endpoint> [<ai_search_index_name>]")
    sys.exit(1)

storage_account_name = sys.argv[1]
blob_container_name = sys.argv[2]
ai_search_endpoint = sys.argv[3]
ai_search_index_name = sys.argv[4] if len(sys.argv) > 4 else "sample-dataset-index"
if not ai_search_endpoint.__contains__("search.windows.net"):
    ai_search_endpoint = f"https://{ai_search_endpoint}.search.windows.net"

class ChainedDeveloperCredential:
    def __init__(self):
        self._credentials = [AzureCliCredential(), InteractiveBrowserCredential()]

    def get_token(self, *scopes, **kwargs):
        last_error = None
        for credential in self._credentials:
            try:
                return credential.get_token(*scopes, **kwargs)
            except Exception as exc:
                last_error = exc
        raise last_error

credential = ChainedDeveloperCredential()

try:
    blob_service_client = BlobServiceClient(account_url=f"https://{storage_account_name}.blob.core.windows.net", credential=credential)
    container_client = blob_service_client.get_container_client(blob_container_name)
    print("Fetching files in container...")
    blob_list = list(container_client.list_blobs())
except Exception as e:
    print(f"Error fetching files: {e}")
    sys.exit(1)

success_count = 0
fail_count = 0
data_list = []

try:
    index_fields = [ 
        SimpleField(name="id", type=SearchFieldDataType.String, key=True),
        SearchableField(name="content", type=SearchFieldDataType.String, searchable=True),
        SearchableField(name="title", type=SearchFieldDataType.String, searchable=True, filterable=True)
    ]
    index = SearchIndex(name=ai_search_index_name, fields=index_fields)

    print("Creating or updating Azure Search index...")
    search_index_client = SearchIndexClient(endpoint=ai_search_endpoint, credential=credential)
    search_index_client.create_or_update_index(index=index)
    print(f"Index '{ai_search_index_name}' created or updated successfully.")
except Exception as e:
    print(f"Error creating/updating index: {e}")
    sys.exit(1)

for idx, blob in enumerate(blob_list, start=1):
    #if blob.name.endswith(".csv"):
    title = blob.name.replace(".csv", "")
    title = title.replace(".json", "")
    title = title.replace(".pdf", "")
    title = title.replace(".docx", "")
    title = title.replace(".pptx", "")
    data = container_client.download_blob(blob.name).readall()
    
    try:
        print(f"Reading data from blob: {blob.name}...")
        #text = data.decode('utf-8')
        # Check if this is a PDF file and process accordingly
        if blob.name.lower().endswith('.pdf'):
            text = extract_pdf_text(data)
        elif blob.name.lower().endswith('.docx'):
            text = extract_docx_text(data)
        else:
            # Original processing for non-PDF files
            text = data.decode('utf-8')
        data_list.append({
            "content": text,
            "id": str(idx),
            "title": title
        })
        success_count += 1
    except Exception as e:
        print(f"Error reading file - {blob.name}: {e}")
        fail_count += 1
        continue

if not data_list:
    print(f"No data to upload to Azure Search index. Success: {success_count}, Failed: {fail_count}")
    sys.exit(1)

try:
    print("Uploading documents to the index...")
    search_client = SearchClient(endpoint=ai_search_endpoint, index_name=ai_search_index_name, credential=credential)
    result = search_client.upload_documents(documents=data_list)
    successes = sum(1 for r in result if getattr(r, "succeeded", False))
    failures = len(data_list) - successes
    print(f"Uploaded documents. Requested: {len(data_list)}, Succeeded: {successes}, Failed: {failures}")
except Exception as e:
    print(f"Error uploading documents: {e}")
    sys.exit(1)

print(f"Processing complete. Success: {success_count}, Failed: {fail_count}")